import streamlit as st
import datetime
from io import BytesIO
import sqlite3
import json
import os
import base64
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
import pandas as pd

# Try importing python-docx for Word file generation
try:
    import docx
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

st.set_page_config(
    page_title="Aplikasi Permohonan Isbat Nikah",
    page_icon="⚖️",
    layout="wide"
)

def apply_custom_theme():
    custom_css = """
    <style>
        /* Mengatur latar belakang aplikasi agar lembut di mata (Off-White/Cream) */
        .stApp {
            background-color: #FAFAFA;
        }
        
        /* Kustomisasi Judul Utama (H1) */
        h1 {
            color: #0B4A2D !important; /* Hijau Tua Pengadilan */
            text-align: center;
            font-family: 'Georgia', serif;
            text-shadow: 1px 1px 2px rgba(0,0,0,0.1);
        }
        
        /* Kustomisasi Sub-judul (H2, H3) */
        h2, h3 {
            color: #115E59 !important; /* Hijau Zamrud */
            border-bottom: 2px solid #D4AF37; /* Garis bawah Kuning Emas */
            padding-bottom: 5px;
            font-family: 'Georgia', serif;
        }

        /* Kustomisasi Kotak Form (Expander) */
        .streamlit-expanderHeader {
            background-color: #EBF3EE !important; /* Hijau sangat muda */
            color: #0B4A2D !important;
            font-weight: 600 !important;
            border-radius: 5px;
            border-left: 5px solid #D4AF37 !important; /* Aksen emas di sisi kiri */
            font-size: 1.1rem;
        }
        
        /* Kustomisasi Tombol Utama (Primary) - Tombol Simpan */
        button[kind="primary"] {
            background-color: #0B4A2D !important;
            color: #D4AF37 !important;
            border: 1px solid #D4AF37 !important;
            border-radius: 8px !important;
            font-weight: bold !important;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
            transition: all 0.3s ease;
        }
        button[kind="primary"]:hover {
            background-color: #D4AF37 !important;
            color: #0B4A2D !important;
            border: 1px solid #0B4A2D !important;
            box-shadow: 0 6px 8px rgba(0,0,0,0.15);
            transform: translateY(-2px);
        }

        /* Kustomisasi Tombol Sekunder (Download) */
        button[kind="secondary"] {
            background-color: #ffffff !important;
            color: #0B4A2D !important;
            border: 1px solid #0B4A2D !important;
            border-radius: 8px !important;
            font-weight: 600 !important;
            transition: all 0.3s ease;
        }
        button[kind="secondary"]:hover {
            background-color: #0B4A2D !important;
            color: #ffffff !important;
            border: 1px solid #D4AF37 !important;
        }
        
        /* Kotak Notifikasi / Info */
        .stAlert {
            background-color: #F0FDF4;
            border-left: 4px solid #16A34A;
            color: #14532D;
            border-radius: 5px;
        }

        /* Garis Pemisah (Divider) mewah */
        hr {
            border: 0;
            height: 2px;
            background-image: linear-gradient(to right, transparent, #D4AF37, transparent);
        }
        
        /* Kustomisasi Header Aplikasi */
        .header-banner {
            background-color: #0B4A2D;
            padding: 20px;
            border-radius: 10px;
            border-bottom: 4px solid #D4AF37;
            margin-bottom: 20px;
            box-shadow: 0 4px 8px rgba(0,0,0,0.1);
            display: flex;
            align-items: center;
            justify-content: center;
            flex-wrap: wrap; /* Menjaga agar tetap rapi di layar HP */
            gap: 25px;
        }
        .banner-logo {
            height: 115px; /* Sedikit dibesarkan agar proporsional */
            width: auto;
            border-radius: 50%; /* Trik CSS memotong kotak putih menjadi bentuk oval */
            border: 2.5px solid #D4AF37; /* Bingkai emas menutupi sisa pinggiran putih */
            box-shadow: 0 4px 12px rgba(0,0,0,0.6); /* Efek bayangan tegas */
            object-fit: cover;
        }
        .header-text-container {
            display: flex;
            flex-direction: column;
            justify-content: center;
            text-align: left;
        }
        @media (max-width: 600px) {
            .header-text-container {
                text-align: center;
                align-items: center;
            }
        }
        .header-title {
            color: #FFD700 !important; /* Kuning Emas Terang memaksa override warna H1 */
            font-size: 2.2rem;
            font-weight: bold;
            margin: 0;
            font-family: 'Georgia', serif;
            text-shadow: 2px 2px 4px rgba(0,0,0,0.3); /* Efek bayangan gelap agar teks lebih menonjol (pop out) */
        }
        .header-subtitle {
            color: #FFFFFF !important; /* Putih bersih agar sangat kontras dibaca */
            font-size: 1.1rem;
            margin-top: 5px;
            font-style: italic;
            text-shadow: 1px 1px 2px rgba(0,0,0,0.2);
        }
    </style>
    """
    st.markdown(custom_css, unsafe_allow_html=True)

# Panggil fungsi tema
apply_custom_theme()

INDONESIAN_MONTHS = [
    "Januari", "Februari", "Maret", "April", "Mei", "Juni",
    "Juli", "Agustus", "September", "Oktober", "November", "Desember"
]

def format_indo_date(dt):
    if not dt: return ""
    return f"{dt.day} {INDONESIAN_MONTHS[dt.month - 1]} {dt.year}"

def calculate_age(birth_date, target_date):
    if not birth_date or not target_date: return 0
    age = target_date.year - birth_date.year
    if (target_date.month, target_date.day) < (birth_date.month, birth_date.day):
        age -= 1
    return max(0, age)

def init_db():
    conn = sqlite3.connect("isbat_nikah.db")
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS permohonan (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            waktu_input TEXT,
            nama_p1 TEXT,
            nik_p1 TEXT,
            nama_p2 TEXT,
            nik_p2 TEXT,
            data_json TEXT
        )
    ''')
    conn.commit()
    conn.close()

def save_to_db(data):
    try:
        conn = sqlite3.connect("isbat_nikah.db")
        c = conn.cursor()
        waktu_input = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        data_json = json.dumps(data, default=str)
        c.execute('''
            INSERT INTO permohonan (waktu_input, nama_p1, nik_p1, nama_p2, nik_p2, data_json)
            VALUES (?, ?, ?, ?, ?, ?)
        ''', (waktu_input, data['p1']['nama'], data['p1']['nik'], data['p2']['nama'], data['p2']['nik'], data_json))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        return False

def load_data_db():
    try:
        conn = sqlite3.connect("isbat_nikah.db")
        c = conn.cursor()
        c.execute("SELECT * FROM permohonan ORDER BY id DESC")
        rows = c.fetchall()
        conn.close()
        return rows
    except:
        return []

init_db()

def build_letter_text(data):
    p1 = data['p1']
    p2 = data['p2']
    tgl_surat_str = format_indo_date(data['tgl_permohonan'])
    
    # Text Anak
    anak_str = "belum dikaruniai anak"
    if data['status_anak'] == 'sudah' and data['anak_list']:
        anak_str = f"telah dikaruniai {len(data['anak_list'])} orang anak, yaitu:\n"
        for idx, child in enumerate(data['anak_list'], 1):
            c_ttl = f"{child['tempat']}, {format_indo_date(child['tgl_lahir'])}"
            anak_str += f"\t{idx}. {child['nama']}, lahir di {c_ttl} (umur {child['umur']} tahun)\n"
        anak_str = anak_str.strip()

    wali_str = data['hubungan_wali']
    if data['hubungan_wali'] != 'Ayah Kandung' and data['alasan_wali']:
        wali_str += f" {data['alasan_wali']}"

    status_p1_full = p1['status'] + (f" ({p1['detail_status_text']})" if p1['detail_status_text'] else "")
    status_p2_full = p2['status'] + (f" ({p2['detail_status_text']})" if p2['detail_status_text'] else "")

    text = f"""Hal : Permohonan Isbat Nikah\t\t\t\tPurwokerto, {tgl_surat_str}

Kepada
Yth. Ketua Pengadilan Agama Purwokerto
di
Purwokerto.

Assalamu Alaikum Wr. Wb.

Dengan hormat,
Yang bertanda tangan di bawah ini :

Nama\t\t\t: {p1['nama']}
NIK\t\t\t: {p1['nik']}
Tempat Tgl Lahir\t: {p1['tempat_lahir']}, {format_indo_date(p1['tgl_lahir'])} (umur {p1['umur']} tahun)
Agama\t\t\t: Islam
Pekerjaan\t\t: {p1['pekerjaan']}
Pendidikan\t\t: {p1['pendidikan']}
Nomor telepon\t\t: {p1['telepon']}
Email\t\t\t: {p1['email']}
Alamat\t\t\t: {p1['alamat']}, selanjutnya disebut sebagai Pemohon I;

Nama\t\t\t: {p2['nama']}
NIK\t\t\t: {p2['nik']}
Tempat Tgl Lahir\t: {p2['tempat_lahir']}, {format_indo_date(p2['tgl_lahir'])} (umur {p2['umur']} tahun)
Agama\t\t\t: Islam
Pekerjaan\t\t: {p2['pekerjaan']}
Pendidikan\t\t: {p2['pendidikan']}
Nomor telepon\t\t: {p2['telepon']}
Email\t\t\t: {p2['email']}
Alamat\t\t\t: {p2['alamat']}, selanjutnya disebut sebagai Pemohon II;

Dengan ini mengajukan pemohonan pengesahan nikah, dengan alasan sebagai berikut:

1. \tBahwa Pemohon I dan Pemohon II telah menikah menurut agama Islam pada tanggal {format_indo_date(data['tgl_nikah'])} di {data['tempat_nikah']} dengan wali nikah adalah {wali_str} bernama {data['nama_wali']}, yang dinikahkan oleh {data['yang_menikahkan']}, dengan maskawin berupa {data['mahar']} yang dibayar tunai, dan dihadiri oleh dua orang saksi masing-masing bernama {data['saksi1']} dan {data['saksi2']};
2. \tBahwa saat menikah Pemohon I berstatus {status_p1_full} dan Pemohon II berstatus {status_p2_full};
3. \tBahwa antara Pemohon I dan Pemohon II tidak ada hubungan keluarga, baik sedarah maupun sesusuan serta Pemohon II juga tidak dalam pinangan laki-laki lain;
4. \tBahwa, pernikahan Pemohon I dengan Pemohon II telah dilaksanakan menurut hukum Islam serta tidak ada masyarakat yang menggugat atau yang meragukan keabsahan atau keberatan atas pernikahan Pemohon I dengan Pemohon II tersebut;
5. \tBahwa dari pernikahan tersebut, Pemohon I dan Pemohon II {anak_str};
6. \tBahwa antara Pemohon I dengan Pemohon II belum pernah terjadi perceraian;
7. \tBahwa, sampai sekarang Pemohon I dengan Pemohon II belum memiliki buku nikah, karena pernikahan Pemohon I dengan Pemohon II tidak terdaftar di Kantor Urusan Agama setempat;
8. \t{data['alasan_tidak_mencatatkan']}
9. \tBahwa, Pemohon I tidak mempunyai isteri yang lain, selain pemohon II;
10. \tBahwa, sekarang Pemohon I dan Pemohon II sangat membutuhkan bukti pernikahan tersebut, untuk mengurus {data['alasan_mohon']};
11. \tBahwa Pemohon I dan Pemohon II bersedia menanggung segala biaya yang ditimbulkan dari pengajuan perkara ini;

Bahwa berdasarkan alasan-alasan tersebut di atas para Pemohon mohon kepada Ketua Pengadilan Agama Purwokerto cq. Majelis hakim yang memeriksa perkara ini berkenan menetapkan sebagai berikut :

Primer :
1. \tMengabulkan permohonan para Pemohon;
2. \tMenyatakan sah perkawinan antara Pemohon I {p1['nama']} dengan Pemohon II, {p2['nama']} yang dilaksanakan pada tanggal {format_indo_date(data['tgl_nikah'])} di {data['tempat_nikah']};
3. \tMenetapkan biaya perkara menurut ketentuan hukum dan perundang-undangan yang berlaku;

Subsider :
- \tAtau bilamana majelis hakim yang memeriksa perkara ini berpendapat lain, mohon penetapan yang seadil-adilnya;

Demikian permohonan para Pemohon, dan atas terkabulnya para Pemohon ucapkan terima kasih.
Wassalam

                                                                                     
           Pemohon I,                                               Pemohon II,      
                                                                                     
                                                                                     
                                                                                     
         ( {p1['nama']} )                                         ( {p2['nama']} )   
"""
    return text

def generate_docx(data):
    if not DOCX_AVAILABLE: return None
    doc = docx.Document()
    
    # Page setup (Atas 3cm, Kiri 4cm, Kanan 2cm, Bawah 3cm)
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(3)
        section.left_margin = Cm(4)
        section.right_margin = Cm(2)

    # Set Default Font to Arial 12pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    p1 = data['p1']
    p2 = data['p2']
    tgl_surat_str = format_indo_date(data['tgl_permohonan'])

    p_header = doc.add_paragraph()
    p_header.add_run(f"Hal : Permohonan Isbat Nikah\t\t\tPurwokerto, {tgl_surat_str}")

    doc.add_paragraph("\tKepada\nYth. Ketua Pengadilan Agama Purwokerto\ndi\nPurwokerto.\n")
    doc.add_paragraph("Assalamu Alaikum Wr. Wb.")
    doc.add_paragraph("Dengan hormat,\nYang bertanda tangan di bawah ini :")

    # Format Biodata dengan model Paragraf (Hanging Indent dan Tab Stops)
    def add_bio_p(label, value):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Mengatur agar baris kedua (jika teks sangat panjang) sejajar di posisi 4.5 cm
        p.paragraph_format.left_indent = Cm(4.5)
        # Menarik kembali label (baris pertama) agar menempel ke margin paling kiri (0 cm)
        p.paragraph_format.first_line_indent = Cm(-4.5)
        
        # Kunci posisi Tab secara absolut (seperti mengatur ruler di Word)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(4.0)) # Tombol Tab 1: Menuju posisi 4.0 cm (untuk Titik Dua)
        tab_stops.add_tab_stop(Cm(4.5)) # Tombol Tab 2: Menuju posisi 4.5 cm (untuk Awal Teks)
        
        # Susunan: [Label] -> Tab -> [:] -> Tab -> [Isi Data]
        p.add_run(f"{label}\t:\t{value}")

    add_bio_p("Nama", p1['nama'])
    add_bio_p("NIK", p1['nik'])
    add_bio_p("Tempat Tgl Lahir", f"{p1['tempat_lahir']}, {format_indo_date(p1['tgl_lahir'])} (umur {p1['umur']} tahun)")
    add_bio_p("Agama", "Islam")
    add_bio_p("Pekerjaan", p1['pekerjaan'])
    add_bio_p("Pendidikan", p1['pendidikan'])
    add_bio_p("Nomor telepon", p1['telepon'])
    add_bio_p("Email", p1['email'])
    add_bio_p("Alamat", f"{p1['alamat']}, selanjutnya disebut sebagai Pemohon I;")
    
    doc.add_paragraph() # Memberikan spasi jarak antara Pemohon I dan Pemohon II
    
    add_bio_p("Nama", p2['nama'])
    add_bio_p("NIK", p2['nik'])
    add_bio_p("Tempat Tgl Lahir", f"{p2['tempat_lahir']}, {format_indo_date(p2['tgl_lahir'])} (umur {p2['umur']} tahun)")
    add_bio_p("Agama", "Islam")
    add_bio_p("Pekerjaan", p2['pekerjaan'])
    add_bio_p("Pendidikan", p2['pendidikan'])
    add_bio_p("Nomor telepon", p2['telepon'])
    add_bio_p("Email", p2['email'])
    add_bio_p("Alamat", f"{p2['alamat']}, selanjutnya disebut sebagai Pemohon II;")

    doc.add_paragraph("\nDengan ini mengajukan pemohonan pengesahan nikah, dengan alasan sebagai berikut:")

    wali_str = data['hubungan_wali']
    if data['hubungan_wali'] != 'Ayah Kandung' and data['alasan_wali']:
        wali_str += f" {data['alasan_wali']}"

    status_p1_full = p1['status'] + (f" ({p1['detail_status_text']})" if p1['detail_status_text'] else "")
    status_p2_full = p2['status'] + (f" ({p2['detail_status_text']})" if p2['detail_status_text'] else "")

    anak_str = "belum dikaruniai anak"
    if data['status_anak'] == 'sudah' and data['anak_list']:
        anak_str = f"telah dikaruniai {len(data['anak_list'])} orang anak, yaitu:\n"
        for idx_child, child in enumerate(data['anak_list'], 1):
            c_ttl = f"{child['tempat']}, {format_indo_date(child['tgl_lahir'])}"
            # Menggunakan \t agar sejajar dengan pengaturan margin indent di bawah
            anak_str += f"{idx_child}.\t{child['nama']}, lahir di {c_ttl} (umur {child['umur']} tahun)\n"
        anak_str = anak_str.strip()

    posita_list = [
        f"Bahwa Pemohon I dan Pemohon II telah menikah menurut agama Islam pada tanggal {format_indo_date(data['tgl_nikah'])} di {data['tempat_nikah']} dengan wali nikah adalah {wali_str} bernama {data['nama_wali']}, yang dinikahkan oleh {data['yang_menikahkan']}, dengan maskawin berupa {data['mahar']} yang dibayar tunai, dan dihadiri oleh dua orang saksi masing-masing bernama {data['saksi1']} dan {data['saksi2']};",
        f"Bahwa saat menikah Pemohon I berstatus {status_p1_full} dan Pemohon II berstatus {status_p2_full};",
        f"Bahwa antara Pemohon I dan Pemohon II tidak ada hubungan keluarga, baik sedarah maupun sesusuan serta Pemohon II juga tidak dalam pinangan laki-laki lain;",
        f"Bahwa, pernikahan Pemohon I dengan Pemohon II telah dilaksanakan menurut hukum Islam serta tidak ada masyarakat yang menggugat atau yang meragukan keabsahan atau keberatan atas pernikahan Pemohon I dengan Pemohon II tersebut;",
        f"Bahwa dari pernikahan tersebut, Pemohon I dan Pemohon II {anak_str};",
        f"Bahwa antara Pemohon I dengan Pemohon II belum pernah terjadi perceraian;",
        f"Bahwa, sampai sekarang Pemohon I dengan Pemohon II belum memiliki buku nikah, karena pernikahan Pemohon I dengan Pemohon II tidak terdaftar di Kantor Urusan Agama setempat;",
        f"{data['alasan_tidak_mencatatkan']}",
        f"Bahwa, Pemohon I tidak mempunyai isteri yang lain, selain pemohon II;",
        f"Bahwa, sekarang Pemohon I dan Pemohon II sangat membutuhkan bukti pernikahan tersebut, untuk mengurus {data['alasan_mohon']};",
        f"Bahwa Pemohon I dan Pemohon II bersedia menanggung segala biaya yang ditimbulkan dari pengajuan perkara ini;"
    ]

    for idx, text in enumerate(posita_list, 1):
        lines = text.split('\n')
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Membuat pola Hanging Indent (Kondisi Baris 1 menjorok ke luar, Baris 2 menjorok ke dalam)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-1.25)
        
        # Angka, lalu pencet Tab (\t), lalu teks
        p.add_run(f"{idx}.\t{lines[0]}")
        
        # Jika ada sub-list (seperti list nama anak yang turun ke bawah)
        if len(lines) > 1:
            for line in lines[1:]:
                p_child = doc.add_paragraph()
                p_child.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Indent digeser lebih dalam lagi sejauh 2.5 cm khusus untuk list nama anak
                p_child.paragraph_format.left_indent = Cm(2.5)
                p_child.paragraph_format.first_line_indent = Cm(-1.25)
                p_child.add_run(line)

    doc.add_paragraph("\nBahwa berdasarkan alasan-alasan tersebut di atas para Pemohon mohon kepada Ketua Pengadilan Agama Purwokerto cq. Majelis hakim yang memeriksa perkara ini berkenan menetapkan sebagai berikut :")
    
    doc.add_paragraph("Primer :")
    petitum_primer = [
        "Mengabulkan permohonan para Pemohon;",
        f"Menyatakan sah perkawinan antara Pemohon I {p1['nama']} dengan Pemohon II, {p2['nama']} yang dilaksanakan pada tanggal {format_indo_date(data['tgl_nikah'])} di {data['tempat_nikah']};",
        "Menetapkan biaya perkara menurut ketentuan hukum dan perundang-undangan yang berlaku;"
    ]
    for idx, text in enumerate(petitum_primer, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-1.25)
        p.add_run(f"{idx}.\t{text}")

    doc.add_paragraph("Subsider :")
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_sub.paragraph_format.left_indent = Cm(1.25)
    p_sub.paragraph_format.first_line_indent = Cm(-1.25)
    p_sub.add_run("-\tAtau bilamana majelis hakim yang memeriksa perkara ini berpendapat lain, mohon penetapan yang seadil-adilnya;")

    doc.add_paragraph("\nDemikian permohonan para Pemohon, dan atas terkabulnya para Pemohon ucapkan terima kasih.\nWassalam\n")

    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    
    cell_p1 = table.cell(0, 0)
    p_p1 = cell_p1.paragraphs[0]
    p_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_p1.add_run("Pemohon I,\n\n\n\n")
    p_p1.add_run(f"( {p1['nama']} )")
    
    cell_p2 = table.cell(0, 1)
    p_p2 = cell_p2.paragraphs[0]
    p_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_p2.add_run("Pemohon II,\n\n\n\n")
    p_p2.add_run(f"( {p2['nama']} )")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Banner Kustom untuk Header Aplikasi ---
logo_path = "logo_PA_Pwt.jpg"
logo_html = '<span style="font-size: 4rem;">⚖️</span>' # Fallback jika gambar tidak ditemukan

if os.path.exists(logo_path):
    with open(logo_path, "rb") as image_file:
        encoded_string = base64.b64encode(image_file.read()).decode()
    
    # Gunakan image/jpeg karena file sekarang bernama logo_PA_Pwt.jpg
    logo_html = f'<img src="data:image/jpeg;base64,{encoded_string}" class="banner-logo" alt="Logo PA Purwokerto">'

st.markdown(f"""
    <div class="header-banner">
        {logo_html}
        <div class="header-text-container">
            <h1 class="header-title">Aplikasi Permohonan Isbat Nikah</h1>
            <p class="header-subtitle">Aplikasi Pembantu Pembuatan Permohonan Isbat Nikah</p>
        </div>
    </div>
""", unsafe_allow_html=True)

if not DOCX_AVAILABLE:
    st.warning("⚠️ Module `python-docx` tidak terdeteksi. Jalankan `pip install python-docx` untuk mengaktifkan ekspor Word.")

col_form, col_preview = st.columns([7, 5])

with col_form:
    st.header("📝 Formulir Permohonan")
    
    with st.expander("1. Informasi Permohonan", expanded=True):
        tgl_permohonan = st.date_input("Tanggal Surat Permohonan", value=datetime.date.today())

    with st.expander("2. Data Pemohon I (Suami)", expanded=True):
        c1, c2 = st.columns(2)
        nama_p1 = c1.text_input("Nama Lengkap Pemohon I", value="Ahmad Bin Fulan")
        nik_p1 = c2.text_input("NIK Pemohon I", value="3302123456780001")
        
        c3, c4 = st.columns(2)
        tempat_lahir_p1 = c3.text_input("Tempat Lahir (Pemohon I)", value="Banyumas")
        tgl_lahir_p1 = c4.date_input("Tanggal Lahir (Pemohon I)", value=datetime.date(1990, 1, 1), min_value=datetime.date(1900, 1, 1), max_value=datetime.date.today())
        
        umur_p1 = calculate_age(tgl_lahir_p1, tgl_permohonan)
        st.info(f"💡 Umur Pemohon I saat ini: **{umur_p1} tahun**")

        pilihan_job_p1 = st.selectbox("Pekerjaan Pemohon I", ["Pegawai BUMN/BUMD", "ASN", "Anggota Polri", "Anggota TNI", "Wiraswasta", "Buruh Harian Lepas", "Petani/Pekebun", "Lain-lain"])
        pekerjaan_p1 = pilihan_job_p1 if pilihan_job_p1 != "Lain-lain" else st.text_input("Sebutkan Pekerjaan Pemohon I", value="")

        pendidikan_p1 = st.selectbox("Pendidikan Pemohon I", ["Tidak Sekolah", "TK", "SD", "SLTP", "SLTA", "D1", "D2", "D3", "D4", "S1", "S2", "S3"], index=4)
        status_p1 = st.selectbox("Status Saat Nikah (Pemohon I)", ["Jejaka", "Duda Cerai", "Duda Mati"])

        detail_status_p1_text = ""
        tgl_ac = None
        tgl_mati = None
        if status_p1 == "Duda Cerai":
            no_ac = st.text_input("No. Akta Cerai (P1)", value="1234/AC/2020/PA.Pwt")
            tgl_ac = st.date_input("Tanggal Akta Cerai (P1)", value=datetime.date(2020, 5, 10))
            pa_penerbit = st.text_input("PA Penerbit (P1)", value="PA Purwokerto")
            detail_status_p1_text = f"Akta Cerai No: {no_ac} tgl {format_indo_date(tgl_ac)} dari {pa_penerbit}"
        elif status_p1 == "Duda Mati":
            tgl_mati = st.date_input("Tanggal Kematian Istri (P1)", value=datetime.date(2019, 1, 10))
            no_surat_m = st.text_input("No. Surat Kematian (P1)", value="474.3/01/2019")
            detail_status_p1_text = f"Srt. Kematian No: {no_surat_m} meninggal tgl {format_indo_date(tgl_mati)}"

        st.markdown("---")
        c5, c6 = st.columns(2)
        telepon_p1 = c5.text_input("Nomor Telepon Pemohon I", value="081234567890")
        email_p1 = c6.text_input("Email Pemohon I", value="ahmad@email.com")
        alamat_p1 = st.text_area("Alamat Lengkap Pemohon I", value="Jl. Merdeka No. 1, RT 01 RW 02, Kec. Purwokerto, Kab. Banyumas")

    with st.expander("3. Data Pemohon II (Istri)", expanded=True):
        c1, c2 = st.columns(2)
        nama_p2 = c1.text_input("Nama Lengkap Pemohon II", value="Siti Bintan")
        nik_p2 = c2.text_input("NIK Pemohon II", value="3302123456780002")
        
        c3, c4 = st.columns(2)
        tempat_lahir_p2 = c3.text_input("Tempat Lahir (Pemohon II)", value="Banyumas")
        tgl_lahir_p2 = c4.date_input("Tanggal Lahir (Pemohon II)", value=datetime.date(1995, 1, 1), min_value=datetime.date(1900, 1, 1), max_value=datetime.date.today())
        
        umur_p2 = calculate_age(tgl_lahir_p2, tgl_permohonan)
        st.info(f"💡 Umur Pemohon II saat ini: **{umur_p2} tahun**")

        pilihan_job_p2 = st.selectbox("Pekerjaan Pemohon II", ["Mengurus Rumah Tangga", "Pegawai BUMN/BUMD", "ASN", "Wiraswasta", "Buruh Harian Lepas", "Lain-lain"])
        pekerjaan_p2 = pilihan_job_p2 if pilihan_job_p2 != "Lain-lain" else st.text_input("Sebutkan Pekerjaan Pemohon II", value="")

        pendidikan_p2 = st.selectbox("Pendidikan Pemohon II", ["Tidak Sekolah", "TK", "SD", "SLTP", "SLTA", "D1", "D2", "D3", "D4", "S1", "S2", "S3"], index=4)
        status_p2 = st.selectbox("Status Saat Nikah (Pemohon II)", ["Perawan", "Janda Cerai", "Janda Mati"])

        detail_status_p2_text = ""
        tgl_ac2 = None
        tgl_mati2 = None
        if status_p2 == "Janda Cerai":
            no_ac2 = st.text_input("No. Akta Cerai (P2)", value="5678/AC/2021/PA.Pwt")
            tgl_ac2 = st.date_input("Tanggal Akta Cerai (P2)", value=datetime.date(2021, 8, 15))
            pa_penerbit2 = st.text_input("PA Penerbit (P2)", value="PA Purwokerto")
            detail_status_p2_text = f"Akta Cerai No: {no_ac2} tgl {format_indo_date(tgl_ac2)} dari {pa_penerbit2}"
        elif status_p2 == "Janda Mati":
            tgl_mati2 = st.date_input("Tanggal Kematian Suami (P2)", value=datetime.date(2020, 2, 10))
            no_surat_m2 = st.text_input("No. Surat Kematian (P2)", value="474.3/02/2020")
            detail_status_p2_text = f"Srt. Kematian No: {no_surat_m2} meninggal tgl {format_indo_date(tgl_mati2)}"
            
        st.markdown("---")
        c5, c6 = st.columns(2)
        telepon_p2 = c5.text_input("Nomor Telepon Pemohon II", value="081298765432")
        email_p2 = c6.text_input("Email Pemohon II", value="siti@email.com")
        alamat_p2 = st.text_area("Alamat Lengkap Pemohon II", value="Jl. Merdeka No. 1, RT 01 RW 02, Kec. Purwokerto, Kab. Banyumas")

    with st.expander("4. Detail Pernikahan Sirri", expanded=True):
        c1, c2 = st.columns(2)
        tgl_nikah = c1.date_input("Tanggal Nikah Sirri", value=datetime.date(2015, 1, 10), min_value=datetime.date(1900, 1, 1), max_value=datetime.date.today())
        tempat_nikah = c2.text_input("Tempat Nikah Sirri", value="Purwokerto")

        hubungan_wali = st.selectbox("Hubungan Wali Nikah dengan P2", ["Ayah Kandung", "Saudara Kandung", "Kakek Kandung (Ayah dari Ayah)", "Paman Kandung (Saudara Ayah)", "Tidak Ada Hubungan Apapun/Hakim"])
        alasan_wali = ""
        if hubungan_wali != "Ayah Kandung":
            alasan_wali = st.text_input("Alasan Wali Nikah (Bukan Ayah)", value="karena ayah kandung sudah meninggal dunia pada tahun 2010")

        nama_wali = st.text_input("Nama Wali Nikah", value="Bpk. Abdullah")
        yang_menikahkan = st.text_input("Nama Kiai/Tokoh yang Menikahkan", value="KH. Ahmad Sholeh")
        mahar = st.text_input("Mas Kawin / Mahar", value="Uang tunai Rp 500.000,- dan seperangkat alat shalat")
        
        c3, c4 = st.columns(2)
        saksi1 = c3.text_input("Saksi Nikah I", value="Saksi 1")
        saksi2 = c4.text_input("Saksi Nikah II", value="Saksi 2")

    with st.expander("5. Alasan & Tujuan Permohonan", expanded=True):
        status_anak = st.selectbox("Keterangan Dikaruniai Anak", ["belum", "sudah"], index=1, format_func=lambda x: "Telah dikaruniai anak" if x == "sudah" else "Belum / Tidak dikaruniai anak")
        
        anak_list = []
        if status_anak == "sudah":
            jumlah_anak = st.number_input("Jumlah Anak", min_value=1, max_value=10, value=2)
            for i in range(int(jumlah_anak)):
                st.markdown(f"**Data Anak ke-{i+1}**")
                c1, c2, c3 = st.columns([4, 3, 3])
                c_nama = c1.text_input(f"Nama Anak #{i+1}", key=f"c_nama_{i}")
                c_tempat = c2.text_input(f"Tempat Lahir #{i+1}", value="Purwokerto", key=f"c_tempat_{i}")
                c_tgl = c3.date_input(f"Tanggal Lahir #{i+1}", value=datetime.date(2020, 1, 1), min_value=datetime.date(1900, 1, 1), max_value=datetime.date.today(), key=f"c_tgl_{i}")
                c_umur = calculate_age(c_tgl, tgl_permohonan)
                anak_list.append({"nama": c_nama, "tempat": c_tempat, "tgl_lahir": c_tgl, "umur": c_umur})

        alasan_tidak_mencatatkan = st.selectbox(
            "Alasan Tidak Mencatatkan Nikah",
            [
                "Bahwa Pemohon I dan Pemohon II telah melaporkan pernikahannya kepada kayim untuk didaftarkan pada Kantor Urusan Agama, namun kayim tersebut tidak melanjutkan pendaftarannya ke Pembantu Pegawai Pencatat Nikah Kantor Urusan Agama;",
                "Bahwa Pemohon I dan Pemohon II telah melaporkan pernikahannya ke Pembantu Pegawai Pencatat Nikah setempat, namun Pembantu Pegawai Pencatat Nikah tersebut tidak melaporkan pencatatan pernikahan tersebut ke Kantor Urusan Agama;",
                "Bahwa Pemohon I dan Pemohon II tidak mendaftarkan pernikahannya ke KUA karena pertimbangan keterbatasan biaya pada saat itu;"
            ], index=2
        )

        pil_alasan_mohon = st.selectbox("Maksud Permohonan", ["penerbitan akta nikah Para Pemohon serta keperluan lainnya", "mengurus akta kelahiran anak Para Pemohon serta keperluan lainnya", "keperluan tersendiri"], index=1)
        alasan_mohon = pil_alasan_mohon if pil_alasan_mohon != "keperluan tersendiri" else st.text_input("Sebutkan Keperluan Anda", value="pendaftaran haji")

form_data = {
    "tgl_permohonan": tgl_permohonan,
    "p1": {"nama": nama_p1, "nik": nik_p1, "tempat_lahir": tempat_lahir_p1, "tgl_lahir": tgl_lahir_p1, "umur": umur_p1, "pekerjaan": pekerjaan_p1, "pendidikan": pendidikan_p1, "status": status_p1, "detail_status_text": detail_status_p1_text, "telepon": telepon_p1, "email": email_p1, "alamat": alamat_p1},
    "p2": {"nama": nama_p2, "nik": nik_p2, "tempat_lahir": tempat_lahir_p2, "tgl_lahir": tgl_lahir_p2, "umur": umur_p2, "pekerjaan": pekerjaan_p2, "pendidikan": pendidikan_p2, "status": status_p2, "detail_status_text": detail_status_p2_text, "telepon": telepon_p2, "email": email_p2, "alamat": alamat_p2},
    "tgl_nikah": tgl_nikah, "tempat_nikah": tempat_nikah, "hubungan_wali": hubungan_wali, "alasan_wali": alasan_wali, "nama_wali": nama_wali, "yang_menikahkan": yang_menikahkan, "mahar": mahar, "saksi1": saksi1, "saksi2": saksi2,
    "status_anak": status_anak, "anak_list": anak_list, "alasan_tidak_mencatatkan": alasan_tidak_mencatatkan, "alasan_mohon": alasan_mohon
}

with col_preview:
    st.header("📄 Pratinjau Dokumen")
    
    # --- LOGIKA VALIDASI POTENSI GAGAL ---
    warnings = []
    
    # Cek Validasi P1 (Suami)
    if status_p1 == "Duda Cerai" and tgl_ac and tgl_nikah < tgl_ac:
        warnings.append("Tanggal nikah sirri lebih dahulu daripada tanggal akta cerai Pemohon I.")
    if status_p1 == "Duda Mati" and tgl_mati and tgl_nikah < tgl_mati:
        warnings.append("Tanggal nikah sirri lebih dahulu daripada tanggal kematian istri Pemohon I.")
        
    # Cek Validasi P2 (Istri)
    if status_p2 == "Janda Cerai" and tgl_ac2:
        if tgl_nikah < tgl_ac2:
            warnings.append("Tanggal nikah sirri lebih dahulu daripada tanggal akta cerai Pemohon II.")
        else:
            selisih_hari = (tgl_nikah - tgl_ac2).days
            if selisih_hari < 90:
                warnings.append(f"Jarak antara tanggal nikah sirri dengan tanggal akta cerai Pemohon II kurang dari 90 hari (Masa iddah berpotensi belum selesai, selisih hanya {selisih_hari} hari).")
    
    if status_p2 == "Janda Mati" and tgl_mati2 and tgl_nikah < tgl_mati2:
        warnings.append("Tanggal nikah sirri lebih dahulu daripada tanggal kematian suami Pemohon II.")

    # Menampilkan Notifikasi Error jika ada peringatan
    if warnings:
        pesan_error = "⚠️ **PERINGATAN: Permohonan berpotensi tidak diloloskan verifikasi karena:**\n"
        for w in warnings:
            pesan_error += f"\n- {w}"
        st.error(pesan_error)
    # ------------------------------------

    if st.button("💾 Simpan Data ke Database", type="primary", use_container_width=True):
        if save_to_db(form_data):
            st.success(f"Berhasil menyimpan data atas nama {nama_p1} & {nama_p2} ke database!")
        else:
            st.error("Gagal menyimpan data ke database.")
            
    st.write("---")
    letter_text = build_letter_text(form_data)
    st.text_area("Hasil Draf Surat Permohonan", value=letter_text, height=580)

    st.subheader("📥 Unduh Dokumen")
    if DOCX_AVAILABLE:
        docx_file = generate_docx(form_data)
        st.download_button(
            label="Unduh Dokumen Word (.docx)",
            data=docx_file,
            file_name=f"Permohonan_Isbat_Nikah_{nama_p1.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="secondary"
        )
    
    st.download_button(
        label="Unduh Teks Draf (.txt)",
        data=letter_text,
        file_name=f"Permohonan_Isbat_Nikah_{nama_p1.replace(' ', '_')}.txt",
        mime="text/plain",
        use_container_width=True,
        type="secondary"
    )

st.write("---")

st.header("🗄️ Panel Admin & Database")

password_input = st.text_input("Masukkan Kata Sandi Admin untuk melihat database:", type="password")

if password_input == "rahasia123":
    st.success("Akses Diberikan!")
    
    with st.expander("Klik untuk melihat histori data yang telah dimasukkan ke database", expanded=True):
        db_data = load_data_db()
        if db_data:
            df_data = []
            for row in db_data:
                base_dict = {"ID": row[0], "Waktu Input": row[1]}
                if len(row) > 6 and row[6]:
                    try:
                        full_data = json.loads(row[6])
                        p1, p2 = full_data.get("p1", {}), full_data.get("p2", {})
                        base_dict.update({
                            "P1 - Nama": p1.get("nama", row[2]), "P1 - NIK": p1.get("nik", row[3]),
                            "P1 - TTL": f"{p1.get('tempat_lahir', '')}, {p1.get('tgl_lahir', '')}",
                            "P1 - Pekerjaan": p1.get("pekerjaan", ""), "P1 - Status": p1.get("status", ""),
                            "P1 - Telepon": p1.get("telepon", ""), "P1 - Email": p1.get("email", ""),
                            "P1 - Alamat": p1.get("alamat", ""),
                            "P2 - Nama": p2.get("nama", row[4]), "P2 - NIK": p2.get("nik", row[5]),
                            "P2 - TTL": f"{p2.get('tempat_lahir', '')}, {p2.get('tgl_lahir', '')}",
                            "P2 - Pekerjaan": p2.get("pekerjaan", ""), "P2 - Status": p2.get("status", ""),
                            "P2 - Telepon": p2.get("telepon", ""), "P2 - Email": p2.get("email", ""),
                            "P2 - Alamat": p2.get("alamat", ""),
                            "Tgl Nikah Sirri": full_data.get("tgl_nikah", ""), "Tempat Nikah": full_data.get("tempat_nikah", ""),
                            "Wali Nikah": full_data.get("nama_wali", ""), "Yg Menikahkan": full_data.get("yang_menikahkan", ""),
                            "Mahar": full_data.get("mahar", ""), "Saksi 1": full_data.get("saksi1", ""), "Saksi 2": full_data.get("saksi2", "")
                        })
                    except:
                        base_dict.update({"P1 - Nama": row[2], "P1 - NIK": row[3], "P2 - Nama": row[4], "P2 - NIK": row[5]})
                
                df_data.append(base_dict)
                
            st.dataframe(df_data, use_container_width=True)
            
            df_export = pd.DataFrame(df_data)
            csv = df_export.to_csv(index=False).encode('utf-8')
            st.download_button(
                label="📊 Unduh Seluruh Data ke Excel (CSV)",
                data=csv,
                file_name=f"Rekap_Data_Isbat_Nikah_{datetime.datetime.now().strftime('%Y%m%d')}.csv",
                mime="text/csv",
                use_container_width=True,
                type="primary"
            )
        else:
            st.info("Belum ada data tersimpan.")
            
    st.write("---")
    st.subheader("📧 Backup Database ke Email")
    st.write("Kirim file `isbat_nikah.db` langsung ke email Anda secara aman.")
    
    with st.form("form_email_backup"):
        default_email = st.secrets.get("email_saya", "")
        default_sandi = st.secrets.get("sandi_aplikasi", "")
        
        email_pengirim = st.text_input("Email Pengirim (Anda)", value=default_email, help="Otomatis terisi dari Secrets jika ada")
        password_app = st.text_input("Sandi Aplikasi Gmail", value=default_sandi, type="password")
        email_penerima = st.text_input("Email Penerima (Tujuan)", value=default_email)
        
        btn_kirim = st.form_submit_button("Kirim File Database (.db) ke Email", type="primary")
        
        if btn_kirim:
            if not email_pengirim or not password_app or not email_penerima:
                st.error("Mohon lengkapi alamat email dan sandi aplikasi!")
            else:
                try:
                    msg = MIMEMultipart()
                    msg['From'] = email_pengirim
                    msg['To'] = email_penerima
                    msg['Subject'] = f"🔒 Backup Database Isbat Nikah - {datetime.datetime.now().strftime('%d %b %Y')}"
                    msg.attach(MIMEText("Terlampir file backup database (isbat_nikah.db) dari Aplikasi Isbat Nikah.", 'plain'))
                    
                    filename = "isbat_nikah.db"
                    if os.path.exists(filename):
                        with open(filename, "rb") as attachment:
                            part = MIMEBase('application', 'octet-stream')
                            part.set_payload(attachment.read())
                            encoders.encode_base64(part)
                            part.add_header('Content-Disposition', f"attachment; filename= {filename}")
                            msg.attach(part)
                            
                        server = smtplib.SMTP('smtp.gmail.com', 587)
                        server.starttls()
                        server.login(email_pengirim, password_app)
                        server.sendmail(email_pengirim, email_penerima, msg.as_string())
                        server.quit()
                        st.success(f"✅ Mantap! File database berhasil dikirim ke {email_penerima}.")
                    else:
                        st.error("Gagal: File database belum terbentuk (belum ada yang menekan tombol Simpan).")
                except Exception as e:
                    st.error(f"Gagal mengirim email. Pastikan sandi aplikasi Anda benar. Error detail: {e}")
elif password_input:
    st.error("❌ Kata sandi salah!")
